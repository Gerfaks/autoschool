from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


def format_report_value(value):
    if value is None or value == '':
        return '-'
    if hasattr(value, 'strftime'):
        return value.strftime('%d.%m.%Y')
    return str(value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    tc_pr.append(shading)


def style_docx_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)

            if row_index == 0:
                set_cell_shading(cell, '243447')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def build_docx_report(title, subtitle, headers, rows):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(10)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run(title)
    heading_run.font.name = 'Arial'
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(36, 52, 71)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(subtitle)
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(90, 101, 117)

    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated_run = generated.add_run(f'Сформировано: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    generated_run.font.name = 'Arial'
    generated_run.font.size = Pt(9)
    generated_run.font.color.rgb = RGBColor(90, 101, 117)

    if not rows:
        empty = document.add_paragraph()
        empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        empty.add_run('Данных для отчета нет.')
    else:
        table = document.add_table(rows=1, cols=len(headers))
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = format_report_value(value)

        style_docx_table(table)

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream


def build_xlsx_report(title, headers, rows, sheet_name='Отчет'):
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = sheet_name[:31]
    worksheet.sheet_view.showGridLines = False

    worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    title_cell = worksheet.cell(row=1, column=1, value=title)
    title_cell.font = Font(name='Arial', size=14, bold=True, color='243447')
    title_cell.alignment = Alignment(horizontal='center')
    worksheet.row_dimensions[1].height = 24

    worksheet.append(headers)
    for cell in worksheet[2]:
        cell.fill = PatternFill('solid', fgColor='E8F1FB')
        cell.font = Font(name='Arial', size=10, bold=True, color='1D2939')
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    worksheet.row_dimensions[2].height = 34

    for row in rows:
        worksheet.append([format_report_value(value) for value in row])

    border = Border(
        left=Side(style='thin', color='D9DEE7'),
        right=Side(style='thin', color='D9DEE7'),
        top=Side(style='thin', color='D9DEE7'),
        bottom=Side(style='thin', color='D9DEE7'),
    )
    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical='top', wrap_text=True)
            if cell.row > 2:
                cell.font = Font(name='Arial', size=10)

    worksheet.freeze_panes = 'A3'
    worksheet.auto_filter.ref = f'A2:{get_column_letter(len(headers))}{worksheet.max_row}'

    for column_index, _header in enumerate(headers, start=1):
        values = [
            format_report_value(worksheet.cell(row=row_index, column=column_index).value)
            for row_index in range(2, worksheet.max_row + 1)
        ]
        width = min(max([len(value) for value in values] + [12]) + 2, 42)
        worksheet.column_dimensions[get_column_letter(column_index)].width = width

    stream = BytesIO()
    workbook.save(stream)
    stream.seek(0)
    return stream
